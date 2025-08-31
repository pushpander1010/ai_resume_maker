import base64
import io
import os
import pickle
import urllib
import uuid
from datetime import date, datetime
from typing import Optional
import pathlib

import mimetypes
from email.message import EmailMessage

import streamlit as st
from dotenv import load_dotenv

from google.oauth2.credentials import Credentials
from google.oauth2 import id_token as google_id_token
from google.auth.transport import requests as google_requests
from google.auth.transport.requests import Request
from google_auth_oauthlib.flow import Flow
from oauthlib.oauth2.rfc6749.errors import InvalidGrantError
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload

from PyPDF2 import PdfReader

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser, PydanticOutputParser
from langchain_google_genai import GoogleGenerativeAI
from langchain_groq import ChatGroq
from langchain_perplexity import ChatPerplexity

from models import ModelState, Details, JD, GmailMessage, Question, QuestionList

# Load env (local file support)
load_dotenv(dotenv_path="environ.env")


# ---- per-user token storage ----
def _tokens_dir() -> str:
    """Return the secure token directory (env TOKENS_DIR or ./tokens)."""
    d = os.environ.get("TOKENS_DIR") or "tokens"
    p = pathlib.Path(d)
    p.mkdir(exist_ok=True)
    # Best-effort restrictive permissions
    try:
        os.chmod(str(p), 0o700)
    except Exception:
        pass
    return str(p)


def _user_token_path(user_sub: str) -> str:
    return os.path.join(_tokens_dir(), f"{user_sub}.pickle")


def _load_creds_for_user(user_sub: str) -> Optional[Credentials]:
    p = _user_token_path(user_sub)
    if not os.path.exists(p):
        return None
    try:
        # Best-effort restrict file perms
        try:
            os.chmod(p, 0o600)
        except Exception:
            pass
        creds = pickle.load(open(p, "rb"))
    except Exception:
        return None
    if creds and creds.expired and creds.refresh_token:
        try:
            creds.refresh(Request())
            pickle.dump(creds, open(p, "wb"))
            try:
                os.chmod(p, 0o600)
            except Exception:
                pass
        except Exception:
            return None
    return creds if creds and creds.valid else None


def _save_creds_for_user(creds: Credentials, client_id: str) -> tuple[str, str]:
    if not creds.id_token:
        creds.refresh(Request())
    claims = google_id_token.verify_oauth2_token(
        creds.id_token, google_requests.Request(), client_id
    )
    user_sub = claims["sub"]
    email = claims.get("email", "")
    name = claims.get("name") or email

    token_path = _user_token_path(user_sub)
    pickle.dump(creds, open(token_path, "wb"))
    try:
        os.chmod(token_path, 0o600)
    except Exception:
        pass
    st.session_state.user_sub = user_sub
    st.session_state.user_email = email
    st.session_state.user_name = name
    return user_sub, email


def _assert_creds_match_session(creds: Credentials, client_id: str):
    """Ensure Google creds belong to the same user as in session_state.
    If not, force sign-out to prevent cross-user leakage.
    """
    try:
        if not creds.id_token:
            creds.refresh(Request())
        claims = google_id_token.verify_oauth2_token(creds.id_token, google_requests.Request(), client_id)
        sub = claims.get("sub")
        email = claims.get("email")
        sess_sub = st.session_state.get("user_sub")
        sess_email = st.session_state.get("user_email")
        if sess_sub and sub and sess_sub != sub:
            st.error("Session mismatch detected. Please sign in again.")
            sign_out()
        if sess_email and email and sess_email.lower() != email.lower():
            st.error("Email mismatch detected. Please sign in again.")
            sign_out()
    except Exception:
        # On verification failure, require re-auth
        st.error("Could not verify credentials. Please sign in again.")
        sign_out()


def _clear_query_params():
    try:
        st.query_params.clear()
    except Exception:
        st.experimental_set_query_params()


def sign_out():
    sub = st.session_state.get("user_sub")
    if sub:
        try:
            p = _user_token_path(sub)
            if os.path.exists(p):
                os.remove(p)
        except Exception:
            pass
    for k in ("user_sub", "user_email", "user_name", "oauth_code_exchanged"):
        st.session_state.pop(k, None)
    st.rerun()


def ensure_google_creds(scopes: list[str], *, force_refresh: bool = False) -> Credentials:
    client_id = os.environ.get("GOOGLE_CLIENT_ID")
    client_secret = os.environ.get("GOOGLE_CLIENT_SECRET")
    redirect_uri = os.environ.get("GOOGLE_REDIRECT_URI")

    if not client_id or not client_secret:
        st.error("Missing GOOGLE_CLIENT_ID / GOOGLE_CLIENT_SECRET.")
        st.stop()
    if not redirect_uri or not redirect_uri.endswith("/"):
        st.error("Set GOOGLE_REDIRECT_URI to your app URL with trailing slash.")
        st.stop()
    if "streamlit.app" in redirect_uri and not redirect_uri.startswith("https://"):
        st.error("On Streamlit Cloud, GOOGLE_REDIRECT_URI must begin with https://")
        st.stop()

    if redirect_uri.startswith("http://"):
        os.environ.setdefault("OAUTHLIB_INSECURE_TRANSPORT", "1")
        os.environ.setdefault("OAUTHLIB_RELAX_TOKEN_SCOPE", "1")

    st.session_state.setdefault("oauth_code_exchanged", False)
    st.session_state.setdefault("user_sub", None)

    if st.session_state.user_sub and not force_refresh:
        cached = _load_creds_for_user(st.session_state.user_sub)
        if cached:
            if not st.session_state.get("user_email"):
                try:
                    if not cached.id_token:
                        cached.refresh(Request())
                    claims = google_id_token.verify_oauth2_token(
                        cached.id_token, google_requests.Request(), client_id
                    )
                    st.session_state.user_email = claims.get("email")
                    st.session_state.user_name = claims.get("name") or claims.get("email")
                except Exception:
                    pass
            # Ensure the cached creds match the current session user
            _assert_creds_match_session(cached, client_id)
            return cached

    try:
        qp = dict(st.query_params)
    except Exception:
        qp = st.experimental_get_query_params()

    def _one(k):
        v = qp.get(k)
        return v[0] if isinstance(v, list) else v

    code = _one("code")
    error = _one("error")
    if error:
        st.error(f"OAuth error: {error}")
        _clear_query_params()
        st.stop()

    client_config = {
        "web": {
            "client_id": client_id,
            "client_secret": client_secret,
            "auth_uri": "https://accounts.google.com/o/oauth2/v2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
            "redirect_uris": [redirect_uri],
            "javascript_origins": [redirect_uri.rstrip("/")],
        }
    }

    flow = Flow.from_client_config(client_config, scopes=scopes)
    flow.redirect_uri = redirect_uri

    if code and not st.session_state.oauth_code_exchanged:
        query = urllib.parse.urlencode(
            {k: (v[0] if isinstance(v, list) else v) for k, v in qp.items()},
            doseq=True,
        )
        authorization_response = f"{redirect_uri}?{query}"
        try:
            flow.fetch_token(authorization_response=authorization_response)
        except InvalidGrantError:
            _clear_query_params()
            st.session_state.oauth_code_exchanged = False
            auth_url, _ = flow.authorization_url(
                access_type="offline",
                include_granted_scopes="true",
                prompt="consent select_account",
            )
            st.link_button("Continue with Google", auth_url)
            st.stop()
        except Exception as e:
            st.error(
                "Failed to fetch token. Ensure this exact redirect URI "
                f"is in your OAuth client:\n{redirect_uri}\n\nDetails: {e}"
            )
            _clear_query_params()
            st.stop()

        creds = flow.credentials
        try:
            _save_creds_for_user(creds, client_id)
        except Exception as e:
            st.error(f"Could not save user token: {e}")
            _clear_query_params()
            st.stop()

        st.session_state.oauth_code_exchanged = True
        _clear_query_params()
        st.rerun()

    if code and st.session_state.oauth_code_exchanged:
        st.stop()

    auth_url, _ = flow.authorization_url(
        access_type="offline",
        include_granted_scopes="true",
        prompt="consent select_account",
    )
    st.link_button("Continue with Google", auth_url)
    st.stop()


def get_model_instance(model_key):
    if not isinstance(model_key, str):
        return model_key
    if model_key.startswith("google|"):
        model_id = model_key.split("|")[1]
        return GoogleGenerativeAI(model=model_id, temperature=0.7)
    elif model_key.startswith("groq|"):
        model_id = model_key.split("|")[1]
        return ChatGroq(model=model_id, temperature=0.7)
    elif model_key.startswith("perplexity|"):
        model_id = model_key.split("|")[1]
        return ChatPerplexity(model=model_id, temperature=0.7)
    else:
        raise ValueError(f"Unknown model: {model_key}")


def passthrough(state: ModelState) -> ModelState:
    return state


def write_email(state: ModelState) -> ModelState:
    parser = PydanticOutputParser(pydantic_object=GmailMessage)
    prompt = PromptTemplate(
        template=(
            "You are an expert email drafter, known for your ability to draft professional emails.\n\n"
            "Given candidate details:\n{candidate_details}\n\n"
            "Draft a professional email based on the job description:\n{jd}\n\n"
            "Required fields:\n"
            "  `to`: string\n"
            "  `subject`: string\n"
            "  `body`: string\n\n"
            "Return the output in STRICT format:\n{template}"
        ),
        input_variables=["candidate_details", "jd"],
        partial_variables={"template": parser.get_format_instructions()},
    )
    chain = prompt | get_model_instance(model_key=state.model) | parser
    output = chain.invoke({"candidate_details": state.candidate_details, "jd": state.jd})
    return {"gmail_message": output}


def convert_docx_to_pdf(state: ModelState) -> ModelState:
    if not state.docx_file or not os.path.exists(state.docx_file):
        return {"pdf_file": None}

    # Try local conversion first (no OAuth, faster)
    try:
        from docx2pdf import convert as _docx2pdf_convert
        base, _ = os.path.splitext(state.docx_file)
        output_path = f"{base}.pdf"
        _docx2pdf_convert(state.docx_file, output_path)
        return {"pdf_file": output_path, "gmail_auth_creds": state.gmail_auth_creds}
    except Exception:
        pass

    SCOPES = [
        "https://www.googleapis.com/auth/drive.file",
        "https://www.googleapis.com/auth/drive.metadata.readonly",
        "https://www.googleapis.com/auth/gmail.send",
        "https://www.googleapis.com/auth/gmail.compose",
        "https://www.googleapis.com/auth/gmail.readonly",
    ]

    st.session_state.oauth_pending_action = "convert_docx_to_pdf"
    st.session_state.oauth_payload = {"docx_file": state.docx_file}

    creds = state.gmail_auth_creds or ensure_google_creds(SCOPES)
    # Verify creds belong to this session user
    try:
        client_id = os.environ.get("GOOGLE_CLIENT_ID") or ""
        if client_id:
            _assert_creds_match_session(creds, client_id)
    except Exception:
        return {}

    if st.session_state.oauth_pending_action == "convert_docx_to_pdf":
        st.session_state.oauth_pending_action = None
        payload = st.session_state.oauth_payload or {}
        st.session_state.oauth_payload = {}

        input_path = payload.get("docx_file") or state.docx_file
        base, _ = os.path.splitext(input_path)
        output_path = f"{base}.pdf"

        drive = build("drive", "v3", credentials=creds)

        file_metadata = {
            "name": os.path.basename(input_path),
            "mimeType": "application/vnd.google-apps.document",
        }
        media = MediaFileUpload(
            input_path,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            resumable=True,
        )
        uploaded = drive.files().create(body=file_metadata, media_body=media, fields="id").execute()
        file_id = uploaded["id"]

        try:
            request = drive.files().export_media(fileId=file_id, mimeType="application/pdf")
            fh = io.BytesIO()
            downloader = MediaIoBaseDownload(fh, request)
            done = False
            while not done:
                _, done = downloader.next_chunk()
            with open(output_path, "wb") as f:
                f.write(fh.getvalue())
            return {"pdf_file": output_path, "gmail_auth_creds": creds}
        finally:
            try:
                drive.files().delete(fileId=file_id).execute()
            except Exception:
                pass

    return {}


def create_draft_with_gmail_auth(state: ModelState) -> ModelState:
    SCOPES = [
        "https://www.googleapis.com/auth/gmail.send",
        "https://www.googleapis.com/auth/gmail.compose",
        "https://www.googleapis.com/auth/gmail.readonly",
    ]
    creds = state.gmail_auth_creds or ensure_google_creds(SCOPES)
    service = build("gmail", "v1", credentials=creds)

    msg = EmailMessage()
    body = (state.gmail_message.body if state.gmail_message else None) or "Default message body"
    msg.set_content(body)

    profile = service.users().getProfile(userId="me").execute()
    from_addr = profile["emailAddress"]
    # Ensure the Gmail profile matches the signed-in user
    sess_email = st.session_state.get("user_email")
    if sess_email and sess_email.lower() != str(from_addr).lower():
        st.error("Authenticated Gmail account mismatch. Please sign in again.")
        sign_out()

    to_addr = (state.gmail_message.to if state.gmail_message and state.gmail_message.to else "example@example.com")
    subject = (state.gmail_message.subject if state.gmail_message and state.gmail_message.subject else "AI Test")

    msg["To"] = to_addr
    msg["From"] = from_addr
    msg["Subject"] = subject

    if state.pdf_file and os.path.exists(state.pdf_file):
        ctype, _ = mimetypes.guess_type(state.pdf_file)
        main, sub = (ctype.split("/", 1) if ctype else ("application", "octet-stream"))
        with open(state.pdf_file, "rb") as f:
            msg.add_attachment(f.read(), maintype=main, subtype=sub, filename=os.path.basename(state.pdf_file))

    raw = base64.urlsafe_b64encode(msg.as_bytes()).decode()
    draft = service.users().drafts().create(userId="me", body={"message": {"raw": raw}}).execute()
    print(f"Draft created: ID = {draft['id']}")
    return {"gmail_auth_creds": creds}


def get_jd(state: ModelState) -> ModelState:
    if state.jd and state.jd.raw_jd:
        return state
    jd_text = input("Enter the job description.")
    if len(jd_text) > 0:
        jd = JD.model_construct(raw_jd=jd_text)
        return {"jd": jd}
    return {"jd": None}


def jd_provided(state: ModelState) -> bool:
    return state.jd is not None


def fill_jd(state: ModelState) -> ModelState:
    jd_obj = state.jd
    try:
        content = jd_obj.get("raw_jd") if isinstance(jd_obj, dict) else jd_obj.raw_jd
    except Exception:
        content = None
    parser = PydanticOutputParser(pydantic_object=JD)
    prompt = PromptTemplate(
        template=(
            "You are good at extracting and filling data in a given template.\n"
            "Task is to fill template: \n{template}, based on given content:\n{content}, "
            "return the output in STRICT format :\n{template}"
        ),
        input_variables=["content"],
        partial_variables={"template": parser.get_format_instructions()},
    )
    chain = prompt | get_model_instance(model_key=state.model) | parser
    output = chain.invoke({"content": content})
    return {"jd": output}


def read_pdf(state: ModelState) -> ModelState:
    reader = PdfReader(state.file_path)
    content = "\n".join(page.extract_text() or "" for page in reader.pages)
    return {"thought": content}


def find_missing(state: ModelState):
    template_parser = PydanticOutputParser(pydantic_object=Details)
    prompt = PromptTemplate(
        template="""
You are a resume evaluator.

Given the following extracted resume information:\n
{resume}

And the required structured format:\n
{template}

Identify the specific fields or types of information that are missing from the resume but are required to fully complete the given template.

For example, if work experience dates or job titles are absent, mention them.
ONLY mention missing items.
List the missing items clearly and concisely, separated by commas.
Only return the field names or types of data that are missing.
""",
        input_variables=["resume", "template"],
    )
    chain = prompt | get_model_instance(model_key=state.model) | StrOutputParser()
    output = chain.invoke({"resume": state.thought, "template": template_parser.get_format_instructions()})
    thought = (state.thought or "") + "\n" + output
    return {"thought": thought}


def ask_questions(state: ModelState):
    parser = PydanticOutputParser(pydantic_object=QuestionList)
    prompt = PromptTemplate(
        template="""
You are an AI assistant helping to improve a candidate's resume.
Given the raw resume {resume}
The following information is missing and needs to be collected:
{missing}

For each missing item, ask a clear and relevant question based on the resume to the user to gather that information. Give examples of possible answer in brackets.
Return your output in the following STRICT format:
{format}
""",
        input_variables=["missing", "resume"],
        partial_variables={"format": parser.get_format_instructions()},
    )
    chain = prompt | get_model_instance(model_key=state.model) | parser
    thought = state.thought or ""
    missing = thought.split("\n")[-1]
    previous_thought = "\n".join(thought.split("\n")[:-1])
    questions = chain.invoke({"missing": missing, "resume": previous_thought})
    return {"questions": questions}


def get_answers(state: ModelState):
    for ques in state.questions.questions:
        a = input(ques.question)
        ques.answer = a
    return {"questions": state.questions}


def fill_details(state: ModelState) -> ModelState:
    parser = PydanticOutputParser(pydantic_object=Details)
    prompt = PromptTemplate(
        template=(
            """Given the candidate details :{candidate_data}\n
            Extract details ,fill and return the following STRICT format:\n{format_instructions}"""
        ),
        input_variables=["candidate_data"],
        partial_variables={"format_instructions": parser.get_format_instructions()},
    )
    chain = prompt | get_model_instance(model_key=state.model) | parser
    output = chain.invoke({"candidate_data": state.thought})
    return {"candidate_details": output}


def resume_improvements(state: ModelState) -> ModelState:
    prompt = PromptTemplate(
        template="""You are an expert resume writer and job application optimizer.

Given the following:
- Job Description : {jd}
- Candidate's Resume: {thought}
- and missing data from resume in form of question answers:{questions}

Your task:
- Improve the resume using ONLY the information already present in the resume.
- Enhance grammar, spelling, and sentence flow.
- Use professional and high-impact action verbs.
- Ensure the resume aligns strongly with the job description.
- Focus on ATS optimization (use relevant keywords from the JD).
- DO NOT fabricate or introduce new experiences, skills, or qualifications.
- The final result must look professional and be compelling.
- Ensure the word count remains under 600 words.

Respond ONLY with the improved resume content.
""",
        input_variables=["jd", "thought", "questions"],
    )
    chain = prompt | get_model_instance(model_key=state.model) | StrOutputParser()
    improved_resume = chain.invoke({"jd": state.jd, "thought": state.thought, "questions": state.questions})
    return {"thought": improved_resume}


def make_resume_docx(state: ModelState) -> ModelState:
    def fmt_date(d):
        if not d:
            return ""
        try:
            if isinstance(d, (date, datetime)):
                return d.strftime("%b %Y")
            for fmt in ("%Y-%m-%d", "%Y-%m", "%d-%m-%Y", "%m/%d/%Y", "%b %Y", "%B %Y", "%Y"):
                try:
                    return datetime.strptime(str(d), fmt).strftime("%b %Y")
                except Exception:
                    pass
            return str(d)
        except Exception:
            return str(d)

    def join_clean(items, sep=" • "):
        return sep.join([str(s) for s in items if s])

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)

    def ensure_style(name, base="Normal", size=10.5, bold=False, all_caps=False, color=RGBColor(0, 0, 0)):
        try:
            stl = doc.styles[name]
        except KeyError:
            stl = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            stl.base_style = doc.styles[base]
        stl.font.name = "Calibri"
        stl._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        stl.font.size = Pt(size)
        stl.font.bold = bold
        stl.font.all_caps = all_caps
        stl.font.color.rgb = RGBColor(0, 0, 0) if color is None else color
        stl.paragraph_format.space_before = Pt(0)
        stl.paragraph_format.space_after = Pt(2)
        stl.paragraph_format.line_spacing = 1.05
        return stl

    ensure_style("SectionHeader", size=10, bold=True, all_caps=True, color=RGBColor(45, 45, 45))
    ensure_style("HeaderName", size=20, bold=True, all_caps=True)
    ensure_style("HeaderContact", size=10)
    ensure_style("Tight", size=10.5)

    def set_spacing(p, before=0, after=2, line=1.05):
        pf = p.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = line

    def add_bullet(text: str):
        try:
            _ = doc.styles["List Bullet"]
            p = doc.add_paragraph(text, style="List Bullet")
        except KeyError:
            p = doc.add_paragraph("\u2022 " + str(text), style="Tight")
        pf = p.paragraph_format
        pf.left_indent = Inches(0.2)
        pf.first_line_indent = Inches(0)
        set_spacing(p, before=0, after=2, line=1.05)
        return p

    def add_hyperlink(paragraph, text, url):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "1155CC")
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(color)
        rPr.append(u)
        new_run.append(rPr)
        text_node = OxmlElement("w:t")
        text_node.text = text
        new_run.append(text_node)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return paragraph

    def add_row_2col(left_text, right_text, left_bold=False):
        table = doc.add_table(rows=1, cols=2)
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "nil")
            borders.append(e)
        tblPr.append(borders)
        try:
            table.columns[0].width = Inches(5.8)
            table.columns[1].width = Inches(1.6)
        except Exception:
            pass
        left, right = table.rows[0].cells
        p_left = left.paragraphs[0]
        run_left = p_left.add_run(left_text)
        run_left.bold = left_bold
        set_spacing(p_left, after=0)
        p_right = right.paragraphs[0]
        p_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p_right.add_run(right_text)
        set_spacing(p_right, after=0)
        return table

    def add_rule():
        t = doc.add_table(rows=1, cols=1)
        tbl = t._tbl
        tblPr = tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "nil")
            borders.append(e)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "999999")
        borders.append(bottom)
        tblPr.append(borders)
        t.rows[0].cells[0].paragraphs[0].add_run("")
        set_spacing(t.rows[0].cells[0].paragraphs[0], after=6)

    if getattr(state.candidate_details, "name", None):
        name_para = doc.add_paragraph(style="HeaderName")
        name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        name_para.add_run(state.candidate_details.name.upper())
        set_spacing(name_para, after=2)

    contact_para = doc.add_paragraph(style="HeaderContact")
    contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_bits = []
    if getattr(state.candidate_details, "email", None):
        contact_bits.append(("mailto:" + state.candidate_details.email, state.candidate_details.email))
    if getattr(state.candidate_details, "phone", None):
        contact_bits.append(("tel:" + state.candidate_details.phone, state.candidate_details.phone))
    profiles = getattr(state.candidate_details, "profiles", []) or []
    for prof in profiles:
        url = getattr(prof, "url", None) if hasattr(prof, "url") else str(prof)
        if url:
            contact_bits.append((url, url))
    for i, (url, text) in enumerate(contact_bits):
        add_hyperlink(contact_para, text, url)
        if i < len(contact_bits) - 1:
            contact_para.add_run("  •  ")
    set_spacing(contact_para, after=6)
    add_rule()

    summary = getattr(state.candidate_details, "summary", None)
    if summary:
        p = doc.add_paragraph("Summary", style="SectionHeader")
        p.paragraph_format.keep_with_next = True
        doc.add_paragraph(summary, style="Tight")

    skills = getattr(state.candidate_details, "skills", None)
    if skills:
        p = doc.add_paragraph("Skills", style="SectionHeader")
        p.paragraph_format.keep_with_next = True
        doc.add_paragraph(join_clean(list(map(str, skills))), style="Tight")

    experience = getattr(state.candidate_details, "experience", None)
    if experience:
        p = doc.add_paragraph("Professional Experience", style="SectionHeader")
        p.paragraph_format.keep_with_next = True
        for exp in experience:
            title = getattr(exp, "title", "") or ""
            company = getattr(exp, "company", "") or ""
            location = getattr(exp, "location", "") or ""
            sd = fmt_date(getattr(exp, "start_date", ""))
            ed_raw = getattr(exp, "end_date", None)
            ed = fmt_date(ed_raw) if ed_raw else "Present"
            left_line = join_clean([s for s in [title, company] if s], sep=", ")
            right_line = join_clean([s for s in [location] if s], sep="")
            add_row_2col(left_line, f"{sd} – {ed}", left_bold=True)
            if right_line:
                add_row_2col("", right_line)
            responsibilities = getattr(exp, "responsibilities", None) or []
            for item in responsibilities:
                add_bullet(str(item))

    projects = getattr(state.candidate_details, "projects", None)
    if projects:
        p = doc.add_paragraph("Projects", style="SectionHeader")
        p.paragraph_format.keep_with_next = True
        for proj in projects:
            name = getattr(proj, "name", "") or ""
            techs = getattr(proj, "technologies", None) or []
            descr = getattr(proj, "description", "") or ""
            date_txt = fmt_date(getattr(proj, "date", "")) if getattr(proj, "date", None) else ""
            link = getattr(proj, "link", None)
            header_left = join_clean([name, f"({date_txt})" if date_txt else ""], sep=" ")
            add_row_2col(header_left, "", left_bold=True)
            if descr:
                doc.add_paragraph(descr, style="Tight")
            if techs:
                doc.add_paragraph("Technologies: " + ", ".join(map(str, techs)), style="Tight")
            if link:
                lp = doc.add_paragraph(style="Tight")
                add_hyperlink(lp, str(link), str(link))

    education = getattr(state.candidate_details, "education", None)
    if education:
        p = doc.add_paragraph("Education", style="SectionHeader")
        p.paragraph_format.keep_with_next = True
        for edu in education:
            degree = getattr(edu, "degree", "") or ""
            institute = getattr(edu, "institute", "") or ""
            sd = fmt_date(getattr(edu, "start_date", ""))
            ed_raw = getattr(edu, "end_date", None)
            ed = fmt_date(ed_raw) if ed_raw else "Present"
            left = join_clean([degree, institute], sep=", ")
            add_row_2col(left, f"{sd} – {ed}")
            gpa = getattr(edu, "gpa", None)
            if gpa:
                doc.add_paragraph(f"GPA: {gpa}", style="Tight")
            coursework = getattr(edu, "coursework", None)
            if coursework:
                doc.add_paragraph("Relevant coursework: " + ", ".join(map(str, coursework)), style="Tight")

    certs = getattr(state.candidate_details, "certifications", None)
    if certs:
        p = doc.add_paragraph("Certifications", style="SectionHeader")
        p.paragraph_format.keep_with_next = True
        for cert in certs:
            name = getattr(cert, "name", "") or ""
            issuer = getattr(cert, "issuer", "") or ""
            cdate = fmt_date(getattr(cert, "date", "")) if getattr(cert, "date", None) else ""
            left = join_clean([name, issuer], sep=" – ")
            add_row_2col(left, cdate)

    base = getattr(state, "file_path", None) or "resume.docx"
    root, _ = os.path.splitext(base)
    output_path = root + ".docx"
    doc.save(output_path)
    return {"docx_file": output_path}


# --- Multiple resume formats (fmt1..fmt5) ---
def _make_resume_with_style(state: ModelState, *, style: str) -> ModelState:
    # Delegate to base implementation but tweak a few visual parameters
    # via simple replacement of fonts/margins/headers.
    # Clone details into a temp state if needed; we’ll reuse base builder
    # by temporarily patching globals. Simpler: copy code with variations.

    # Choose style presets
    presets = {
        # Modern: Calibri, neutral gray accent
        "fmt1": {
            "font": "Calibri",
            "size": 10.5,
            "margins": 0.5,
            "header_caps": True,
            "accent": (45, 45, 45),  # dark gray
            "header_banner": False,
            "bullet": "• ",
        },
        # Classic: Times New Roman, minimal accent, hyphen bullets
        "fmt2": {
            "font": "Times New Roman",
            "size": 11,
            "margins": 1.0,
            "header_caps": False,
            "accent": (0, 0, 0),  # black
            "header_banner": False,
            "bullet": "- ",
        },
        # Clean: Arial with blue accent
        "fmt3": {
            "font": "Arial",
            "size": 10.5,
            "margins": 0.75,
            "header_caps": True,
            "accent": (47, 84, 235),  # blue
            "header_banner": False,
            "bullet": "• ",
        },
        # Tight: Verdana with green accent and square bullets
        "fmt4": {
            "font": "Verdana",
            "size": 10,
            "margins": 0.6,
            "header_caps": False,
            "accent": (0, 135, 90),  # green
            "header_banner": False,
            "bullet": "▪ ",
        },
        # Professional: Georgia with maroon accent and header banner
        "fmt5": {
            "font": "Georgia",
            "size": 11,
            "margins": 0.8,
            "header_caps": True,
            "accent": (128, 0, 0),  # maroon
            "header_banner": True,
            "bullet": "• ",
        },
        # Additional expanded styles
        "fmt6": {"font": "Garamond", "size": 11, "margins": 0.8, "header_caps": True, "accent": (47, 84, 235), "header_banner": False, "bullet": "• "},
        "fmt7": {"font": "Cambria", "size": 11, "margins": 0.7, "header_caps": True, "accent": (0, 128, 128), "header_banner": False, "bullet": "– "},
        "fmt8": {"font": "Tahoma", "size": 10.5, "margins": 0.6, "header_caps": False, "accent": (224, 108, 0), "header_banner": False, "bullet": "• "},
        "fmt9": {"font": "Trebuchet MS", "size": 10.5, "margins": 0.6, "header_caps": True, "accent": (102, 51, 153), "header_banner": False, "bullet": "▪ "},
        "fmt10": {"font": "Century Gothic", "size": 10.5, "margins": 0.6, "header_caps": False, "accent": (96, 125, 139), "header_banner": False, "bullet": "• "},
        "fmt11": {"font": "Palatino Linotype", "size": 11, "margins": 0.8, "header_caps": True, "accent": (0, 51, 102), "header_banner": False, "bullet": "– "},
        "fmt12": {"font": "Calibri", "size": 11, "margins": 0.5, "header_caps": True, "accent": (0, 188, 212), "header_banner": True, "bullet": "• "},
        "fmt13": {"font": "Arial", "size": 10, "margins": 0.5, "header_caps": False, "accent": (63, 81, 181), "header_banner": False, "bullet": "◦ "},
        "fmt14": {"font": "Georgia", "size": 10.5, "margins": 0.6, "header_caps": False, "accent": (233, 30, 99), "header_banner": False, "bullet": "• "},
        "fmt15": {"font": "Verdana", "size": 10, "margins": 0.5, "header_caps": True, "accent": (46, 125, 50), "header_banner": True, "bullet": "▪ "},
        "fmt16": {"font": "Times New Roman", "size": 12, "margins": 1.0, "header_caps": False, "accent": (33, 150, 243), "header_banner": False, "bullet": "– "},
        "fmt17": {"font": "Cambria", "size": 10, "margins": 0.5, "header_caps": True, "accent": (0, 105, 92), "header_banner": False, "bullet": "• "},
        "fmt18": {"font": "Garamond", "size": 11, "margins": 0.7, "header_caps": False, "accent": (121, 85, 72), "header_banner": True, "bullet": "▪ "},
        "fmt19": {"font": "Tahoma", "size": 10, "margins": 0.5, "header_caps": True, "accent": (0, 150, 136), "header_banner": False, "bullet": "• "},
        "fmt20": {"font": "Trebuchet MS", "size": 10.5, "margins": 0.7, "header_caps": False, "accent": (255, 87, 34), "header_banner": False, "bullet": "– "},
        "fmt21": {"font": "Century Gothic", "size": 11, "margins": 0.7, "header_caps": True, "accent": (156, 39, 176), "header_banner": True, "bullet": "• "},
        "fmt22": {"font": "Palatino Linotype", "size": 10.5, "margins": 0.6, "header_caps": False, "accent": (205, 220, 57), "header_banner": False, "bullet": "◦ "},
        "fmt23": {"font": "Arial", "size": 11, "margins": 0.75, "header_caps": True, "accent": (0, 0, 0), "header_banner": False, "bullet": "• "},
        "fmt24": {"font": "Georgia", "size": 11, "margins": 0.9, "header_caps": True, "accent": (63, 81, 181), "header_banner": True, "bullet": "▪ "},
        "fmt25": {"font": "Verdana", "size": 10.5, "margins": 0.5, "header_caps": False, "accent": (33, 33, 33), "header_banner": False, "bullet": "• "},
        "fmt26": {"font": "Times New Roman", "size": 11, "margins": 0.5, "header_caps": True, "accent": (76, 175, 80), "header_banner": False, "bullet": "– "},
        "fmt27": {"font": "Cambria", "size": 11, "margins": 0.9, "header_caps": False, "accent": (244, 67, 54), "header_banner": False, "bullet": "▪ "},
        "fmt28": {"font": "Garamond", "size": 10.5, "margins": 0.6, "header_caps": True, "accent": (33, 150, 243), "header_banner": True, "bullet": "• "},
        "fmt29": {"font": "Tahoma", "size": 11, "margins": 0.8, "header_caps": False, "accent": (158, 158, 158), "header_banner": False, "bullet": "◦ "},
        "fmt30": {"font": "Trebuchet MS", "size": 10.5, "margins": 0.6, "header_caps": True, "accent": (121, 85, 72), "header_banner": True, "bullet": "• "},
        
    }
    p = presets.get(style, presets["fmt1"])
    # Align fmt28 with Streamlit preview: red accent, Arial, banner + sidebar
    if style == "fmt28":
        try:
            p.update({
                "font": "Arial",
                "accent": (244, 67, 54),
                "header_banner": True,
            })
            p["sidebar"] = True
        except Exception:
            pass
    if style == "fmt12":
        # Ensure sidebar like preview
        p["sidebar"] = True

    # Build a document similarly to make_resume_docx but with preset tweaks
    from docx import Document  # local import to avoid confusion

    def fmt_date(d):
        if not d:
            return ""
        try:
            if isinstance(d, (date, datetime)):
                return d.strftime("%b %Y")
            for fmt in ("%Y-%m-%d", "%Y-%m", "%d-%m-%Y", "%m/%d/%Y", "%b %Y", "%B %Y", "%Y"):
                try:
                    return datetime.strptime(str(d), fmt).strftime("%b %Y")
                except Exception:
                    pass
            return str(d)
        except Exception:
            return str(d)

    def join_clean(items, sep=" • "):
        return sep.join([str(s) for s in items if s])

    doc = Document()
    for section in doc.sections:
        m = p["margins"]
        section.top_margin = Inches(m)
        section.bottom_margin = Inches(m)
        section.left_margin = Inches(m)
        section.right_margin = Inches(m)

    normal = doc.styles["Normal"]
    normal.font.name = p["font"]
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), p["font"])
    normal.font.size = Pt(p["size"])

    def ensure_style(name, base="Normal", size=None, bold=False, all_caps=False, color=RGBColor(0, 0, 0)):
        try:
            stl = doc.styles[name]
        except KeyError:
            stl = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            stl.base_style = doc.styles[base]
        stl.font.name = p["font"]
        stl._element.rPr.rFonts.set(qn("w:eastAsia"), p["font"])
        stl.font.size = Pt(size if size else p["size"])
        stl.font.bold = bold
        stl.font.all_caps = all_caps
        stl.font.color.rgb = color
        stl.paragraph_format.space_before = Pt(0)
        stl.paragraph_format.space_after = Pt(2)
        stl.paragraph_format.line_spacing = 1.05
        return stl

    accent_rgb = p.get("accent", (45, 45, 45))
    ensure_style(
        "SectionHeader",
        size=max(p["size"] - 0.5, 9.5),
        bold=True,
        all_caps=p["header_caps"],
        color=RGBColor(*accent_rgb),
    )
    ensure_style("HeaderName", size=p["size"] + 9, bold=True, all_caps=p["header_caps"])  # bigger name
    ensure_style("HeaderContact", size=p["size"])
    ensure_style("Tight", size=p["size"])

    def set_spacing(pgh, before=0, after=2, line=1.05):
        pf = pgh.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = line

    def add_bullet(text: str):
        try:
            _ = doc.styles["List Bullet"]
            pgh = doc.add_paragraph(text, style="List Bullet")
        except KeyError:
            pgh = doc.add_paragraph(p.get("bullet", "• ") + str(text), style="Tight")
        pf = pgh.paragraph_format
        pf.left_indent = Inches(0.2)
        pf.first_line_indent = Inches(0)
        set_spacing(pgh, before=0, after=2, line=1.05)
        return pgh

    def add_hyperlink(paragraph, text, url):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "1155CC")
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(color)
        rPr.append(u)
        new_run.append(rPr)
        text_node = OxmlElement("w:t")
        text_node.text = text
        new_run.append(text_node)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return paragraph

    def add_row_2col(left_text, right_text, left_bold=False):
        table = doc.add_table(rows=1, cols=2)
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "nil")
            borders.append(e)
        tblPr.append(borders)
        try:
            table.columns[0].width = Inches(5.8)
            table.columns[1].width = Inches(1.6)
        except Exception:
            pass
        left, right = table.rows[0].cells
        p_left = left.paragraphs[0]
        run_left = p_left.add_run(left_text)
        run_left.bold = left_bold
        set_spacing(p_left, after=0)
        p_right = right.paragraphs[0]
        p_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p_right.add_run(right_text)
        set_spacing(p_right, after=0)
        return table

    def add_rule():
        t = doc.add_table(rows=1, cols=1)
        tbl = t._tbl
        tblPr = tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "nil")
            borders.append(e)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(*accent_rgb))
        borders.append(bottom)
        tblPr.append(borders)
        t.rows[0].cells[0].paragraphs[0].add_run("")
        set_spacing(t.rows[0].cells[0].paragraphs[0], after=6)

    # Header
    def add_header_banner(text: str):
        # Create colored banner with white text
        t = doc.add_table(rows=1, cols=1)
        cell = t.rows[0].cells[0]
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*accent_rgb))
        tcPr.append(shd)
        pgh = cell.paragraphs[0]
        pgh.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = pgh.add_run(text)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        run.font.size = Pt(p["size"] + 10)
        set_spacing(pgh, after=4)

    if getattr(state.candidate_details, "name", None):
        header_text = state.candidate_details.name.upper() if p["header_caps"] else state.candidate_details.name
        if p.get("header_banner"):
            add_header_banner(header_text)
        else:
            name_para = doc.add_paragraph(style="HeaderName")
            name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = name_para.add_run(header_text)
            # Subtle color tint for some styles
            if style in ("fmt3", "fmt4", "fmt5"):
                run.font.color.rgb = RGBColor(*accent_rgb)
            set_spacing(name_para, after=2)

    # Sidebar layout for formats like fmt28
    if p.get("sidebar"):
        table = doc.add_table(rows=1, cols=2)
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "nil")
            borders.append(e)
        tblPr.append(borders)
        try:
            table.columns[0].width = Inches(2.2)
            table.columns[1].width = Inches(5.2)
        except Exception:
            pass
        left_cell, right_cell = table.rows[0].cells

        # Left column: contacts and skills
        lc = left_cell.paragraphs[0]
        lc = left_cell.add_paragraph("Contact", style="SectionHeader")
        set_spacing(lc, after=2)
        contact_bits = []
        if getattr(state.candidate_details, "email", None):
            contact_bits.append(("mailto:" + state.candidate_details.email, state.candidate_details.email))
        if getattr(state.candidate_details, "phone", None):
            contact_bits.append(("tel:" + state.candidate_details.phone, state.candidate_details.phone))
        for url_item in (getattr(state.candidate_details, "profiles", []) or []):
            url = getattr(url_item, "url", None) if hasattr(url_item, "url") else str(url_item)
            if url:
                contact_bits.append((url, url))
        for (url, text) in contact_bits:
            pgh = left_cell.add_paragraph(style="Tight")
            add_hyperlink(pgh, text, url)

        skills = getattr(state.candidate_details, "skills", None)
        if skills:
            p_sk = left_cell.add_paragraph("Skills", style="SectionHeader")
            set_spacing(p_sk, after=2)
            for s in skills:
                left_cell.add_paragraph((p.get("bullet") or "• ") + str(s), style="Tight")

        # Right column: summary and main sections
        summary = getattr(state.candidate_details, "summary", None)
        if summary:
            p_rh = right_cell.add_paragraph("Summary", style="SectionHeader")
            p_rh.paragraph_format.keep_with_next = True
            right_cell.add_paragraph(summary, style="Tight")

        experience = getattr(state.candidate_details, "experience", None)
        if experience:
            p_rh = right_cell.add_paragraph("Professional Experience", style="SectionHeader")
            p_rh.paragraph_format.keep_with_next = True
            for exp in experience:
                title = getattr(exp, "title", "") or ""
                company = getattr(exp, "company", "") or ""
                location = getattr(exp, "location", "") or ""
                sd = fmt_date(getattr(exp, "start_date", ""))
                ed_raw = getattr(exp, "end_date", None)
                ed = fmt_date(ed_raw) if ed_raw else "Present"
                header = f"{', '.join([s for s in [title, company] if s])}    {sd} – {ed}"
                rp = right_cell.add_paragraph(header, style="Tight")
                rp.runs and setattr(rp.runs[0], 'bold', True)
                if location:
                    right_cell.add_paragraph(location, style="Tight")
                for item in getattr(exp, "responsibilities", None) or []:
                    right_cell.add_paragraph((p.get("bullet") or "• ") + str(item), style="Tight")

        projects = getattr(state.candidate_details, "projects", None)
        if projects:
            p_rh = right_cell.add_paragraph("Projects", style="SectionHeader")
            p_rh.paragraph_format.keep_with_next = True
            for proj in projects:
                name = getattr(proj, "name", "") or ""
                descr = getattr(proj, "description", "") or ""
                date_txt = fmt_date(getattr(proj, "date", "")) if getattr(proj, "date", None) else ""
                header_left = ' '.join([x for x in [name, f"({date_txt})" if date_txt else ""] if x])
                rp = right_cell.add_paragraph(header_left, style="Tight")
                rp.runs and setattr(rp.runs[0], 'bold', True)
                if descr:
                    right_cell.add_paragraph(descr, style="Tight")

        education = getattr(state.candidate_details, "education", None)
        if education:
            p_rh = right_cell.add_paragraph("Education", style="SectionHeader")
            p_rh.paragraph_format.keep_with_next = True
            for edu in education:
                degree = getattr(edu, "degree", "") or ""
                institute = getattr(edu, "institute", "") or ""
                sd = fmt_date(getattr(edu, "start_date", ""))
                ed_raw = getattr(edu, "end_date", None)
                ed = fmt_date(ed_raw) if ed_raw else "Present"
                header = f"{', '.join([s for s in [degree, institute] if s])}    {sd} – {ed}"
                right_cell.add_paragraph(header, style="Tight")

        certs = getattr(state.candidate_details, "certifications", None)
        if certs:
            p_rh = right_cell.add_paragraph("Certifications", style="SectionHeader")
            p_rh.paragraph_format.keep_with_next = True
            for cert in certs:
                name = getattr(cert, "name", "") or ""
                issuer = getattr(cert, "issuer", "") or ""
                cdate = fmt_date(getattr(cert, "date", "")) if getattr(cert, "date", None) else ""
                right_cell.add_paragraph(
                    (" – ".join([s for s in [name, issuer] if s])) + (f"  {cdate}" if cdate else ""),
                    style="Tight",
                )

        base = getattr(state, "file_path", None) or "resume.docx"
        root, _ = os.path.splitext(base)
        output_path = f"{root}_{style}.docx"
        doc.save(output_path)
        return {"docx_file": output_path}

    contact_para = doc.add_paragraph(style="HeaderContact")
    contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_bits = []
    if getattr(state.candidate_details, "email", None):
        contact_bits.append(("mailto:" + state.candidate_details.email, state.candidate_details.email))
    if getattr(state.candidate_details, "phone", None):
        contact_bits.append(("tel:" + state.candidate_details.phone, state.candidate_details.phone))
    profiles = getattr(state.candidate_details, "profiles", []) or []
    for prof in profiles:
        url = getattr(prof, "url", None) if hasattr(prof, "url") else str(prof)
        if url:
            contact_bits.append((url, url))
    for i, (url, text) in enumerate(contact_bits):
        add_hyperlink(contact_para, text, url)
        if i < len(contact_bits) - 1:
            contact_para.add_run("  •  ")
    set_spacing(contact_para, after=6)
    add_rule()

    # Summary
    summary = getattr(state.candidate_details, "summary", None)
    if summary:
        pgh = doc.add_paragraph("Summary", style="SectionHeader")
        pgh.paragraph_format.keep_with_next = True
        doc.add_paragraph(summary, style="Tight")

    # Skills
    skills = getattr(state.candidate_details, "skills", None)
    if skills:
        pgh = doc.add_paragraph("Skills", style="SectionHeader")
        pgh.paragraph_format.keep_with_next = True
        doc.add_paragraph(join_clean(list(map(str, skills))), style="Tight")

    # Experience
    experience = getattr(state.candidate_details, "experience", None)
    if experience:
        pgh = doc.add_paragraph("Professional Experience", style="SectionHeader")
        pgh.paragraph_format.keep_with_next = True
        for exp in experience:
            title = getattr(exp, "title", "") or ""
            company = getattr(exp, "company", "") or ""
            location = getattr(exp, "location", "") or ""
            sd = fmt_date(getattr(exp, "start_date", ""))
            ed_raw = getattr(exp, "end_date", None)
            ed = fmt_date(ed_raw) if ed_raw else "Present"
            left_line = join_clean([s for s in [title, company] if s], sep=", ")
            right_line = join_clean([s for s in [location] if s], sep="")
            add_row_2col(left_line, f"{sd} – {ed}", left_bold=True)
            if right_line:
                add_row_2col("", right_line)
            responsibilities = getattr(exp, "responsibilities", None) or []
            for item in responsibilities:
                add_bullet(str(item))

    # Projects
    projects = getattr(state.candidate_details, "projects", None)
    if projects:
        pgh = doc.add_paragraph("Projects", style="SectionHeader")
        pgh.paragraph_format.keep_with_next = True
        for proj in projects:
            name = getattr(proj, "name", "") or ""
            techs = getattr(proj, "technologies", None) or []
            descr = getattr(proj, "description", "") or ""
            date_txt = fmt_date(getattr(proj, "date", "")) if getattr(proj, "date", None) else ""
            link = getattr(proj, "link", None)
            header_left = join_clean([name, f"({date_txt})" if date_txt else ""], sep=" ")
            add_row_2col(header_left, "", left_bold=True)
            if descr:
                doc.add_paragraph(descr, style="Tight")
            if techs:
                doc.add_paragraph("Technologies: " + ", ".join(map(str, techs)), style="Tight")
            if link:
                lp = doc.add_paragraph(style="Tight")
                add_hyperlink(lp, str(link), str(link))

    # Education
    education = getattr(state.candidate_details, "education", None)
    if education:
        pgh = doc.add_paragraph("Education", style="SectionHeader")
        pgh.paragraph_format.keep_with_next = True
        for edu in education:
            degree = getattr(edu, "degree", "") or ""
            institute = getattr(edu, "institute", "") or ""
            sd = fmt_date(getattr(edu, "start_date", ""))
            ed_raw = getattr(edu, "end_date", None)
            ed = fmt_date(ed_raw) if ed_raw else "Present"
            left = join_clean([degree, institute], sep=", ")
            add_row_2col(left, f"{sd} – {ed}")
            gpa = getattr(edu, "gpa", None)
            if gpa:
                doc.add_paragraph(f"GPA: {gpa}", style="Tight")
            coursework = getattr(edu, "coursework", None)
            if coursework:
                doc.add_paragraph("Relevant coursework: " + ", ".join(map(str, coursework)), style="Tight")

    # Certifications
    certs = getattr(state.candidate_details, "certifications", None)
    if certs:
        pgh = doc.add_paragraph("Certifications", style="SectionHeader")
        pgh.paragraph_format.keep_with_next = True
        for cert in certs:
            name = getattr(cert, "name", "") or ""
            issuer = getattr(cert, "issuer", "") or ""
            cdate = fmt_date(getattr(cert, "date", "")) if getattr(cert, "date", None) else ""
            left = join_clean([name, issuer], sep=" – ")
            add_row_2col(left, cdate)

    base = getattr(state, "file_path", None) or "resume.docx"
    root, _ = os.path.splitext(base)
    output_path = f"{root}_{style}.docx"
    doc.save(output_path)
    return {"docx_file": output_path}


def make_resume_docx_1(state: ModelState) -> ModelState:
    return _make_resume_with_style(state, style="fmt1")


def make_resume_docx_2(state: ModelState) -> ModelState:
    return _make_resume_with_style(state, style="fmt2")


def make_resume_docx_3(state: ModelState) -> ModelState:
    return _make_resume_with_style(state, style="fmt3")


def make_resume_docx_4(state: ModelState) -> ModelState:
    return _make_resume_with_style(state, style="fmt4")


def make_resume_docx_5(state: ModelState) -> ModelState:
    return _make_resume_with_style(state, style="fmt5")


def select_resume_format(state: ModelState) -> str:
    # Return routing key used in graph conditional edges
    value = (state.resume_format or "fmt1").lower()
    if value not in {"fmt1", "fmt2", "fmt3", "fmt4", "fmt5"}:
        return "fmt1"
    return value


def is_email_in_jd(state: ModelState):
    email = None
    if state.jd:
        if isinstance(state.jd, dict):
            email = state.jd.get("email")
        else:
            email = getattr(state.jd, "email", None)
    if email and "@" in str(email):
        return "email_present"
    return "email_absent"


def write_referral(state: ModelState):
    prompt = PromptTemplate.from_template(
        """
You are a job applicant seeking a referral. Write a short and professional LinkedIn-style referral message (60–100 words max) to someone working at the company.

Use the following:
- Job Description:
{jd}

- Resume Summary:
{resume}

Write in a polite, concise tone. Don't assume familiarity.
"""
    )

    chain = prompt | get_model_instance(model_key=state.model) | StrOutputParser()
    output = chain.invoke({"jd": state.jd, "resume": state.thought})
    gm = GmailMessage.model_construct(
        to=(state.jd.email if state.jd and getattr(state.jd, "email", None) else None),
        subject="Referral Request",
        body=output,
    )
    return {"referral_message": output, "gmail_message": gm}


FONTS = {
    "calibri": ("Calibri", 10.5),
    "times": ("Times New Roman", 11),
    "arial": ("Arial", 10.5),
    "verdana": ("Verdana", 10),
    "georgia": ("Georgia", 11),
    "garamond": ("Garamond", 11),
    "cambria": ("Cambria", 11),
    "tahoma": ("Tahoma", 10.5),
    "trebuchet": ("Trebuchet MS", 10.5),
    "centurygothic": ("Century Gothic", 10.5),
}

COLORS = {
    "blue": (47, 84, 235),
    "teal": (0, 150, 136),
    "green": (46, 125, 50),
    "red": (244, 67, 54),
    "purple": (156, 39, 176),
    "slate": (96, 125, 139),
    "orange": (255, 87, 34),
    "navy": (0, 51, 102),
    "maroon": (128, 0, 0),
    "gray": (45, 45, 45),
}

LAYOUTS = {
    # classic: single column, neutral caps, no banner/sidebar
    "classic": {"banner": False, "sidebar": False, "margins": 0.75, "header_caps": True},
    # banner: single column with header banner
    "banner": {"banner": True, "sidebar": False, "margins": 0.75, "header_caps": True},
    # sidebar: two-column layout
    "sidebar": {"banner": True, "sidebar": True, "margins": 0.75, "header_caps": True},
    # compact: tighter margins, single column
    "compact": {"banner": False, "sidebar": False, "margins": 0.5, "header_caps": True},
    # modern: single column with tighter margins
    "modern": {"banner": False, "sidebar": False, "margins": 0.6, "header_caps": True},
    # minimal: single column, larger margins, no caps
    "minimal": {"banner": False, "sidebar": False, "margins": 1.0, "header_caps": False},
    # elegant: banner + larger margins
    "elegant": {"banner": True, "sidebar": False, "margins": 1.0, "header_caps": True},
    # sidebar-wide: wider left column for skills/contact
    "sidebar-wide": {"banner": True, "sidebar": True, "margins": 0.75, "header_caps": True, "sidebar_widths": (2.7, 4.7)},
}


def _build_preset_from_state(state: ModelState) -> dict:
    layout_key = (state.resume_layout or "classic").lower()
    font_key = (state.resume_font or "calibri").lower()
    color_key = (state.resume_color or "blue").lower()

    font_name, font_size = FONTS.get(font_key, FONTS["calibri"])
    accent_rgb = COLORS.get(color_key, COLORS["blue"])
    layout = LAYOUTS.get(layout_key, LAYOUTS["classic"]).copy()
    p = {
        "font": font_name,
        "size": font_size,
        "margins": layout.get("margins", 0.75),
        "header_caps": layout.get("header_caps", True),
        "accent": accent_rgb,
        "header_banner": layout.get("banner", False),
        "sidebar": layout.get("sidebar", False),
        "sidebar_widths": layout.get("sidebar_widths"),
        "bullet": "• ",
    }
    return p


def make_resume_docx_styled(state: ModelState) -> ModelState:
    # Legacy support: if resume_format set to fmtXX, keep special cases (e.g., fmt28)
    fmt = (state.resume_format or "").lower()
    if fmt.startswith("fmt"):
        # Map some special legacy codes to new presets
        if fmt == "fmt28":
            state.resume_layout = "sidebar"
            state.resume_font = "arial"
            state.resume_color = "red"
    p = _build_preset_from_state(state)

    # Now render with the generalized preset p
    return _render_doc_with_preset(state, p)


def _render_doc_with_preset(state: ModelState, p: dict) -> ModelState:
    # Similar to _make_resume_with_style but driven solely by p
    from docx import Document

    def fmt_date(d):
        if not d:
            return ""
        try:
            if isinstance(d, (date, datetime)):
                return d.strftime("%b %Y")
            for fmt in ("%Y-%m-%d", "%Y-%m", "%d-%m-%Y", "%m/%d/%Y", "%b %Y", "%B %Y", "%Y"):
                try:
                    return datetime.strptime(str(d), fmt).strftime("%b %Y")
                except Exception:
                    pass
            return str(d)
        except Exception:
            return str(d)

    def join_clean(items, sep=" • "):
        return sep.join([str(s) for s in items if s])

    doc = Document()
    for section in doc.sections:
        m = p.get("margins", 0.75)
        section.top_margin = Inches(m)
        section.bottom_margin = Inches(m)
        section.left_margin = Inches(m)
        section.right_margin = Inches(m)

    normal = doc.styles["Normal"]
    normal.font.name = p["font"]
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), p["font"])
    normal.font.size = Pt(p.get("size", 10.5))

    def ensure_style(name, base="Normal", size=None, bold=False, all_caps=False, color=RGBColor(0, 0, 0)):
        try:
            stl = doc.styles[name]
        except KeyError:
            stl = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            stl.base_style = doc.styles[base]
        stl.font.name = p["font"]
        stl._element.rPr.rFonts.set(qn("w:eastAsia"), p["font"])
        stl.font.size = Pt(size if size else p.get("size", 10.5))
        stl.font.bold = bold
        stl.font.all_caps = all_caps
        stl.font.color.rgb = color
        stl.paragraph_format.space_before = Pt(0)
        stl.paragraph_format.space_after = Pt(2)
        stl.paragraph_format.line_spacing = 1.05
        return stl

    accent_rgb = p.get("accent", (45, 45, 45))
    ensure_style("SectionHeader", size=max(p.get("size", 10.5) - 0.5, 9.5), bold=True, all_caps=p.get("header_caps", True), color=RGBColor(*accent_rgb))
    ensure_style("HeaderName", size=p.get("size", 10.5) + 9, bold=True, all_caps=p.get("header_caps", True))
    ensure_style("HeaderContact", size=p.get("size", 10.5))
    ensure_style("Tight", size=p.get("size", 10.5))

    def set_spacing(pgh, before=0, after=2, line=1.05):
        pf = pgh.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = line

    def add_hyperlink(paragraph, text, url):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "1155CC")
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(color)
        rPr.append(u)
        new_run.append(rPr)
        text_node = OxmlElement("w:t")
        text_node.text = text
        new_run.append(text_node)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return paragraph

    # Header + optional banner
    def add_rule():
        t = doc.add_table(rows=1, cols=1)
        tbl = t._tbl
        tblPr = tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "nil")
            borders.append(e)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(*accent_rgb))
        borders.append(bottom)
        tblPr.append(borders)
        t.rows[0].cells[0].paragraphs[0].add_run("")
        set_spacing(t.rows[0].cells[0].paragraphs[0], after=6)

    if getattr(state.candidate_details, "name", None):
        name_text = state.candidate_details.name.upper() if p.get("header_caps", True) else state.candidate_details.name
        if p.get("header_banner"):
            t = doc.add_table(rows=1, cols=1)
            cell = t.rows[0].cells[0]
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*accent_rgb))
            tcPr.append(shd)
            pgh = cell.paragraphs[0]
            pgh.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = pgh.add_run(name_text)
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(p.get("size", 10.5) + 10)
            set_spacing(pgh, after=4)
        else:
            name_para = doc.add_paragraph(style="HeaderName")
            name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            name_para.add_run(name_text)
            set_spacing(name_para, after=2)

    # Sidebar vs single column
    if p.get("sidebar"):
        table = doc.add_table(rows=1, cols=2)
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "nil")
            borders.append(e)
        tblPr.append(borders)
        try:
            w = p.get("sidebar_widths", (2.2, 5.2))
            table.columns[0].width = Inches(w[0])
            table.columns[1].width = Inches(w[1])
        except Exception:
            pass
        left_cell, right_cell = table.rows[0].cells

        # Left: Contact + Skills + Links
        lc = left_cell.add_paragraph("Contact", style="SectionHeader")
        set_spacing(lc, after=2)
        contact_bits = []
        if getattr(state.candidate_details, "email", None):
            contact_bits.append(("mailto:" + state.candidate_details.email, state.candidate_details.email))
        if getattr(state.candidate_details, "phone", None):
            contact_bits.append(("tel:" + state.candidate_details.phone, state.candidate_details.phone))
        for url_item in (getattr(state.candidate_details, "profiles", []) or []):
            url = getattr(url_item, "url", None) if hasattr(url_item, "url") else str(url_item)
            if url:
                contact_bits.append((url, url))
        for (url, text) in contact_bits:
            pgh = left_cell.add_paragraph(style="Tight")
            add_hyperlink(pgh, text, url)

        skills = getattr(state.candidate_details, "skills", None)
        if skills:
            p_sk = left_cell.add_paragraph("Skills", style="SectionHeader")
            set_spacing(p_sk, after=2)
            for s in skills:
                left_cell.add_paragraph(p.get("bullet", "• ") + str(s), style="Tight")

        # Right content
        summary = getattr(state.candidate_details, "summary", None)
        if summary:
            p_rh = right_cell.add_paragraph("Summary", style="SectionHeader")
            p_rh.paragraph_format.keep_with_next = True
            right_cell.add_paragraph(summary, style="Tight")

        experience = getattr(state.candidate_details, "experience", None)
        if experience:
            p_rh = right_cell.add_paragraph("Professional Experience", style="SectionHeader")
            p_rh.paragraph_format.keep_with_next = True
            for exp in experience:
                title = getattr(exp, "title", "") or ""
                company = getattr(exp, "company", "") or ""
                location = getattr(exp, "location", "") or ""
                sd = fmt_date(getattr(exp, "start_date", ""))
                ed_raw = getattr(exp, "end_date", None)
                ed = fmt_date(ed_raw) if ed_raw else "Present"
                header = ", ".join([s for s in [title, company] if s]) + f"    {sd} – {ed}"
                rp = right_cell.add_paragraph(header, style="Tight")
                rp.runs and setattr(rp.runs[0], 'bold', True)
                if location:
                    right_cell.add_paragraph(location, style="Tight")
                for item in getattr(exp, "responsibilities", None) or []:
                    right_cell.add_paragraph(p.get("bullet", "• ") + str(item), style="Tight")

        projects = getattr(state.candidate_details, "projects", None)
        if projects:
            p_rh = right_cell.add_paragraph("Projects", style="SectionHeader")
            p_rh.paragraph_format.keep_with_next = True
            for proj in projects:
                name = getattr(proj, "name", "") or ""
                descr = getattr(proj, "description", "") or ""
                date_txt = fmt_date(getattr(proj, "date", "")) if getattr(proj, "date", None) else ""
                header_left = ' '.join([x for x in [name, f"({date_txt})" if date_txt else ""] if x])
                rp = right_cell.add_paragraph(header_left, style="Tight")
                rp.runs and setattr(rp.runs[0], 'bold', True)
                if descr:
                    right_cell.add_paragraph(descr, style="Tight")

        education = getattr(state.candidate_details, "education", None)
        if education:
            p_rh = right_cell.add_paragraph("Education", style="SectionHeader")
            p_rh.paragraph_format.keep_with_next = True
            for edu in education:
                degree = getattr(edu, "degree", "") or ""
                institute = getattr(edu, "institute", "") or ""
                sd = fmt_date(getattr(edu, "start_date", ""))
                ed_raw = getattr(edu, "end_date", None)
                ed = fmt_date(ed_raw) if ed_raw else "Present"
                header = ", ".join([s for s in [degree, institute] if s]) + f"    {sd} – {ed}"
                right_cell.add_paragraph(header, style="Tight")

        certs = getattr(state.candidate_details, "certifications", None)
        if certs:
            p_rh = right_cell.add_paragraph("Certifications", style="SectionHeader")
            p_rh.paragraph_format.keep_with_next = True
            for cert in certs:
                name = getattr(cert, "name", "") or ""
                issuer = getattr(cert, "issuer", "") or ""
                cdate = fmt_date(getattr(cert, "date", "")) if getattr(cert, "date", None) else ""
                right_cell.add_paragraph((" – ".join([s for s in [name, issuer] if s])) + (f"  {cdate}" if cdate else ""), style="Tight")

        base = getattr(state, "file_path", None) or "resume.docx"
        root, _ = os.path.splitext(base)
        output_path = f"{root}_{state.resume_layout or 'layout'}_{state.resume_font or 'font'}_{state.resume_color or 'color'}.docx"
        doc.save(output_path)
        return {"docx_file": output_path}

    # Single column flow
    contact_para = doc.add_paragraph(style="HeaderContact")
    contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_bits = []
    if getattr(state.candidate_details, "email", None):
        contact_bits.append(("mailto:" + state.candidate_details.email, state.candidate_details.email))
    if getattr(state.candidate_details, "phone", None):
        contact_bits.append(("tel:" + state.candidate_details.phone, state.candidate_details.phone))
    for prof in getattr(state.candidate_details, "profiles", []) or []:
        url = getattr(prof, "url", None) if hasattr(prof, "url") else str(prof)
        if url:
            contact_bits.append((url, url))
    for i, (url, text) in enumerate(contact_bits):
        add_hyperlink(contact_para, text, url)
        if i < len(contact_bits) - 1:
            contact_para.add_run("  •  ")
    set_spacing(contact_para, after=6)
    add_rule()

    summary = getattr(state.candidate_details, "summary", None)
    if summary:
        pgh = doc.add_paragraph("Summary", style="SectionHeader")
        pgh.paragraph_format.keep_with_next = True
        doc.add_paragraph(summary, style="Tight")

    skills = getattr(state.candidate_details, "skills", None)
    if skills:
        pgh = doc.add_paragraph("Skills", style="SectionHeader")
        pgh.paragraph_format.keep_with_next = True
        doc.add_paragraph(join_clean(list(map(str, skills))), style="Tight")

    experience = getattr(state.candidate_details, "experience", None)
    if experience:
        pgh = doc.add_paragraph("Professional Experience", style="SectionHeader")
        pgh.paragraph_format.keep_with_next = True
        for exp in experience:
            title = getattr(exp, "title", "") or ""
            company = getattr(exp, "company", "") or ""
            location = getattr(exp, "location", "") or ""
            sd = fmt_date(getattr(exp, "start_date", ""))
            ed_raw = getattr(exp, "end_date", None)
            ed = fmt_date(ed_raw) if ed_raw else "Present"
            left_line = ", ".join([s for s in [title, company] if s])
            right_line = location or ""
            tbl = doc.add_table(rows=1, cols=2)
            left, right = tbl.rows[0].cells
            p_left = left.paragraphs[0]
            p_left.add_run(left_line).bold = True
            p_right = right.paragraphs[0]
            p_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            p_right.add_run(f"{sd} – {ed}")
            for item in getattr(exp, "responsibilities", None) or []:
                doc.add_paragraph(p.get("bullet", "• ") + str(item), style="Tight")

    projects = getattr(state.candidate_details, "projects", None)
    if projects:
        pgh = doc.add_paragraph("Projects", style="SectionHeader")
        pgh.paragraph_format.keep_with_next = True
        for proj in projects:
            name = getattr(proj, "name", "") or ""
            descr = getattr(proj, "description", "") or ""
            date_txt = fmt_date(getattr(proj, "date", "")) if getattr(proj, "date", None) else ""
            header_left = ' '.join([x for x in [name, f"({date_txt})" if date_txt else ""] if x])
            rp = doc.add_paragraph(header_left, style="Tight")
            rp.runs and setattr(rp.runs[0], 'bold', True)
            if descr:
                doc.add_paragraph(descr, style="Tight")

    education = getattr(state.candidate_details, "education", None)
    if education:
        pgh = doc.add_paragraph("Education", style="SectionHeader")
        pgh.paragraph_format.keep_with_next = True
        for edu in education:
            degree = getattr(edu, "degree", "") or ""
            institute = getattr(edu, "institute", "") or ""
            sd = fmt_date(getattr(edu, "start_date", ""))
            ed_raw = getattr(edu, "end_date", None)
            ed = fmt_date(ed_raw) if ed_raw else "Present"
            rp = doc.add_paragraph(f"{degree}, {institute}    {sd} – {ed}", style="Tight")
            set_spacing(rp, after=0)

    certs = getattr(state.candidate_details, "certifications", None)
    if certs:
        pgh = doc.add_paragraph("Certifications", style="SectionHeader")
        pgh.paragraph_format.keep_with_next = True
        for cert in certs:
            name = getattr(cert, "name", "") or ""
            issuer = getattr(cert, "issuer", "") or ""
            cdate = fmt_date(getattr(cert, "date", "")) if getattr(cert, "date", None) else ""
            doc.add_paragraph((" – ".join([s for s in [name, issuer] if s])) + (f"  {cdate}" if cdate else ""), style="Tight")

    base = getattr(state, "file_path", None) or "resume.docx"
    root, _ = os.path.splitext(base)
    output_path = f"{root}_{state.resume_layout or 'layout'}_{state.resume_font or 'font'}_{state.resume_color or 'color'}.docx"
    doc.save(output_path)
    return {"docx_file": output_path}
